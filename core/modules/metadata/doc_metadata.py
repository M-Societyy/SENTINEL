# sentinel - extraccion de metadatos de documentos office
# m-society & c1q_

import asyncio
import io
from typing import Optional

import structlog

from modules.base import ModuloBase
from schemas.modulos import ResultadoEnriquecimiento

log = structlog.get_logger()


class DocMetadata(ModuloBase):
    nombre = "doc_metadata"
    categoria = "metadata"
    descripcion = "extraccion de metadatos de documentos docx"

    async def ejecutar(self, objetivo: str, parametros: dict = None) -> ResultadoEnriquecimiento:
        entidades = []
        relaciones = []
        resultados = {"objetivo": objetivo}

        doc_bytes = None
        if objetivo.startswith("http"):
            resp = await self.request_con_rate_limit(objetivo, servicio="default")
            if resp and resp.status_code == 200:
                doc_bytes = resp.content
        else:
            try:
                with open(objetivo, "rb") as f:
                    doc_bytes = f.read()
            except FileNotFoundError:
                return ResultadoEnriquecimiento(
                    fuente=self.nombre, tipo="document",
                    error="archivo no encontrado", confianza=0.0,
                )

        if not doc_bytes:
            return ResultadoEnriquecimiento(
                fuente=self.nombre, tipo="document",
                error="no se pudo obtener el documento", confianza=0.0,
            )

        loop = asyncio.get_event_loop()
        metadata = await loop.run_in_executor(None, self._extraer_metadata, doc_bytes)

        if metadata:
            resultados["metadata"] = metadata

            if metadata.get("autor"):
                entidades.append({
                    "tipo": "person", "valor": metadata["autor"],
                    "datos": {"fuente": "doc_metadata"},
                    "confianza": 0.6,
                })
            if metadata.get("ultimo_modificado_por"):
                entidades.append({
                    "tipo": "person", "valor": metadata["ultimo_modificado_por"],
                    "datos": {"fuente": "doc_metadata", "accion": "ultima_modificacion"},
                    "confianza": 0.5,
                })
            if metadata.get("empresa"):
                entidades.append({
                    "tipo": "organization", "valor": metadata["empresa"],
                    "datos": {"fuente": "doc_metadata"},
                    "confianza": 0.6,
                })

        self._entidades_encontradas = len(entidades)

        return ResultadoEnriquecimiento(
            fuente=self.nombre, tipo="document",
            datos=resultados, confianza=0.7,
            entidades_nuevas=entidades,
            relaciones_nuevas=relaciones,
        )

    def _extraer_metadata(self, doc_bytes: bytes) -> Optional[dict]:
        try:
            from docx import Document
            doc = Document(io.BytesIO(doc_bytes))
            props = doc.core_properties

            return {
                "autor": props.author,
                "titulo": props.title,
                "asunto": props.subject,
                "categoria": props.category,
                "palabras_clave": props.keywords,
                "comentarios": props.comments,
                "ultimo_modificado_por": props.last_modified_by,
                "revision": props.revision,
                "fecha_creacion": str(props.created) if props.created else None,
                "fecha_modificacion": str(props.modified) if props.modified else None,
                "empresa": getattr(props, "company", None),
                "total_parrafos": len(doc.paragraphs),
                "total_tablas": len(doc.tables),
            }
        except Exception as e:
            return {"error": str(e)}
